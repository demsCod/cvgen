from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.models import AdaptationResult, CandidateProfile, ExportPayload


def ensure_output_dir() -> Path:
    base = Path.home() / ".cvgen" / "exports"
    base.mkdir(parents=True, exist_ok=True)
    return base


def export_docx(
    profile: CandidateProfile,
    adaptation: AdaptationResult,
    output_dir: Path,
) -> tuple[Path, Path]:
    doc = Document()
    doc.add_heading(profile.full_name, level=1)
    if profile.summary:
        doc.add_paragraph(profile.summary)

    doc.add_heading("Compétences clés", level=2)
    doc.add_paragraph(", ".join(profile.skills[:12]))

    doc.add_heading("Expériences", level=2)
    for experience in profile.experiences:
        doc.add_heading(f"{experience.role} – {experience.company}", level=3)
        doc.add_paragraph(f"{experience.start_date} - {experience.end_date or 'Présent'}")
        for achievement in experience.achievements:
            doc.add_paragraph(achievement, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("CV adapté", level=1)
    for line in adaptation.adapted_resume.splitlines():
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("Lettre de motivation", level=1)
    for line in adaptation.adapted_cover_letter.splitlines():
        doc.add_paragraph(line)

    resume_doc = output_dir / "cvgen_resume.docx"
    doc.save(resume_doc)

    letter_doc = Document()
    letter_doc.add_heading("Lettre de motivation", level=1)
    for paragraph in adaptation.adapted_cover_letter.split("\n\n"):
        letter_doc.add_paragraph(paragraph)
    letter_path = output_dir / "cvgen_letter.docx"
    letter_doc.save(letter_path)

    return resume_doc, letter_path


def export_pdf(adaptation: AdaptationResult, output_dir: Path) -> Path:
    output_path = output_dir / "cvgen_resume.pdf"
    pdf = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    text_object = pdf.beginText(50, height - 50)
    text_object.setFont("Helvetica-Bold", 16)
    text_object.textLine("CV adapté")
    text_object.moveCursor(0, 20)

    text_object.setFont("Helvetica", 11)
    for line in adaptation.adapted_resume.splitlines():
        text_object.textLine(line)
    text_object.moveCursor(0, 20)

    text_object.setFont("Helvetica-Bold", 16)
    text_object.textLine("Lettre de motivation")
    text_object.moveCursor(0, 20)

    text_object.setFont("Helvetica", 11)
    for line in adaptation.adapted_cover_letter.splitlines():
        text_object.textLine(line)

    pdf.drawText(text_object)
    pdf.showPage()
    pdf.save()
    return output_path


def export_documents(profile: CandidateProfile, adaptation: AdaptationResult, fmt: str) -> ExportPayload:
    output_dir = ensure_output_dir()
    resume_docx, letter_docx = export_docx(profile, adaptation, output_dir)
    resume_pdf = export_pdf(adaptation, output_dir)

    resume_path = resume_pdf if fmt == "pdf" else resume_docx

    return ExportPayload(
        resume_path=str(resume_path),
        cover_letter_path=str(letter_docx),
    )
